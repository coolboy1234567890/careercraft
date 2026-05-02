from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_bcrypt import Bcrypt
from groq import Groq
from dotenv import load_dotenv
import os
import requests
import stripe
from datetime import date
import resend
import secrets


load_dotenv()
stripe.api_key = os.getenv("STRIPE_SECRET_KEY")
resend.api_key = os.getenv("RESEND_API_KEY")
app = Flask(__name__)
CORS(app)
bcrypt = Bcrypt(app)
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

SUPABASE_URL = os.getenv("SUPABASE_URL")
# Use the service role key — this bypasses RLS safely from your server.
# Never expose this key in the frontend. The anon key is only for client-side Supabase auth (which we don't use).
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY") or os.getenv("SUPABASE_KEY")  # fallback keeps dev working
HEADERS = {
    "apikey": SUPABASE_KEY,
    "Authorization": f"Bearer {SUPABASE_KEY}",
    "Content-Type": "application/json"
}

FREE_DAILY_LIMIT = 5

def db_get(table, filters=""):
    res = requests.get(f"{SUPABASE_URL}/rest/v1/{table}?{filters}", headers=HEADERS)
    return res.json()

def db_post(table, data):
    res = requests.post(f"{SUPABASE_URL}/rest/v1/{table}", json=data, headers={**HEADERS, "Prefer": "return=representation"})
    return res.json()

def db_patch(table, filters, data):
    res = requests.patch(f"{SUPABASE_URL}/rest/v1/{table}?{filters}", json=data, headers={**HEADERS, "Prefer": "return=representation"})
    return res.json()

def db_delete(table, filters):
    res = requests.delete(f"{SUPABASE_URL}/rest/v1/{table}?{filters}", headers=HEADERS)
    return res.status_code

@app.route("/delete-resume", methods=["POST"])
def delete_resume():
    data = request.json
    resume_id = data.get("resume_id")
    user_id = data.get("user_id")
    if not resume_id or not user_id:
        return jsonify({"error": "Missing resume_id or user_id"}), 400
    status = db_delete("resumes", f"id=eq.{resume_id}&user_id=eq.{user_id}")
    if status in (200, 204):
        return jsonify({"message": "Resume deleted"})
    return jsonify({"error": "Could not delete resume"}), 500


@app.route("/delete-history", methods=["POST"])
def delete_history():
    data = request.json
    user_id = data.get("user_id")
    if not user_id:
        return jsonify({"error": "Missing user_id"}), 400
    status = db_delete("resumes", f"user_id=eq.{user_id}")
    if status in (200, 204):
        return jsonify({"message": "History cleared"})
    return jsonify({"error": "Could not clear history"}), 500


@app.route("/delete-account", methods=["POST"])
def delete_account():
    data = request.json
    user_id = data.get("user_id")
    password = data.get("password")
    if not user_id or not password:
        return jsonify({"error": "Missing user_id or password"}), 400
    users = db_get("users", f"id=eq.{user_id}&select=id,password")
    if not users:
        return jsonify({"error": "User not found"}), 404
    if not bcrypt.check_password_hash(users[0]["password"], password):
        return jsonify({"error": "Incorrect password"}), 401
    db_delete("resumes", f"user_id=eq.{user_id}")
    db_delete("usage",   f"user_id=eq.{user_id}")
    db_delete("users",   f"id=eq.{user_id}")
    return jsonify({"message": "Account deleted"})


@app.route("/")
def home():
    return jsonify({"message": "CareerCraft API is running! 🚀"})

@app.route("/signup", methods=["POST"])
def signup():
    data = request.json
    email = data.get("email", "").lower().strip()
    password = data.get("password", "")
    promo = data.get("promo_code", "").strip().upper()

    if not email or not password:
        return jsonify({"error": "Email and password are required"}), 400

    existing = db_get("users", f"email=eq.{email}&select=id")
    if existing:
        return jsonify({"error": "An account with this email already exists"}), 400

    valid_promo = os.getenv("PROMO_CODE", "").upper()
    plan = "proplus" if promo and promo == valid_promo else "free"

    hashed = bcrypt.generate_password_hash(password).decode("utf-8")
    user = db_post("users", {"email": email, "password": hashed, "plan": plan})

    if not user or "id" not in user[0]:
        return jsonify({"error": "Failed to create account"}), 500

    u = user[0]
    return jsonify({
        "message": "Account created!",
        "user": {"id": u["id"], "email": u["email"], "plan": u["plan"]}
    })

@app.route("/login", methods=["POST"])
def login():
    data = request.json
    email = data.get("email", "").lower().strip()
    password = data.get("password", "")

    users = db_get("users", f"email=eq.{email}&select=id,email,password,plan")
    if not users:
        return jsonify({"error": "No account found with this email"}), 404

    user = users[0]
    if not bcrypt.check_password_hash(user["password"], password):
        return jsonify({"error": "Incorrect password"}), 401

    return jsonify({"message": "Logged in!", "user": {"id": user["id"], "email": user["email"], "plan": user["plan"]}})

@app.route("/generate", methods=["POST"])
def generate_resume():
    data = request.json
    user_id = data.get("user_id")
    plan = data.get("plan", "free")

    if plan == "free" and user_id:
        today = str(date.today())
        usage = db_get("usage", f"user_id=eq.{user_id}&date=eq.{today}&select=count,id")

        if usage:
            if usage[0]["count"] >= FREE_DAILY_LIMIT:
                return jsonify({"error": "daily_limit", "message": "You've used all 5 free generations today. Upgrade to Pro for unlimited!"}), 429
            db_patch("usage", f"user_id=eq.{user_id}&date=eq.{today}", {"count": usage[0]["count"] + 1})
        else:
            db_post("usage", {"user_id": user_id, "date": today, "count": 1})

    job_description = data.get("job_description", "")
    experience = data.get("experience", "")
    skills = data.get("skills", "")
    education = data.get("education", "")
    tone = data.get("tone", "professional")
    voice_sample = data.get("voice_sample")
    job_title = data.get("job_title", "Position")

    voice_instruction = ""
    if voice_sample and plan in ["pro", "proplus"]:
        voice_instruction = f"""
        The user has provided a writing sample to match their voice:
        ---
        {voice_sample}
        ---
        Analyze their tone, vocabulary, and sentence style.
        Write the resume and cover letter to sound like THEM, not like a generic AI.
        """

    prompt = f"""
    You are an expert resume writer and career coach with 20 years of experience.

    Job Description:
    {job_description}

    Candidate Information:
    - Experience: {experience}
    - Skills: {skills}
    - Education: {education}
    - Preferred tone: {tone}

    {voice_instruction}

    STRICT RULES — follow every one:
    1. NEVER invent experience, certifications, clearances, or degrees not mentioned by the candidate
    2. Only claim skills the candidate actually listed — if a skill matches the job but wasnt mentioned, use bridging language like "experience applicable to X environments"
    3. Add realistic context — include scale, team size, or impact where logical (e.g. "systems serving millions of users")
    4. Include measurable outcomes where possible (e.g. "reduced deployment time by 30%") but only if they are believable given the candidates background
    5. Use natural varied language — avoid buzzword overload and keyword stuffing
    6. Rotate wording naturally — dont repeat the same phrase in every bullet
    7. Make it sound like a real human wrote it, not an AI
    8. Place ATS keywords naturally throughout, not crammed into every line
    9. Connect the candidates real experience to the job using bridging statements
    10. Cover letter must only reference verified experience from the candidates input

    Please write:
    1. A tailored professional resume using keywords from the job description
    2. A personalized cover letter

    Output ONLY the resume and cover letter. No explanations, no notes, no ATS optimization section.
    """

    chat = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )
    result = chat.choices[0].message.content

    if user_id:
        db_post("resumes", {
            "user_id": user_id,
            "content": result,
            "job_title": job_title
        })

    return jsonify({"result": result})

@app.route("/create-checkout", methods=["POST"])
def create_checkout():
    data = request.json
    plan = data.get("plan")
    user_id = data.get("user_id")
    email = data.get("email")

    price_id = os.getenv("STRIPE_PRO_PRICE_ID") if plan == "pro" else os.getenv("STRIPE_PROPLUS_PRICE_ID")

    try:
        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            mode="subscription",
            customer_email=email,
            line_items=[{"price": price_id, "quantity": 1}],
            success_url="http://127.0.0.1:5500/index.html?success=true&plan=" + plan,
            cancel_url="http://127.0.0.1:5500/index.html?canceled=true",
            metadata={"user_id": user_id, "plan": plan}
        )
        return jsonify({"url": session.url})
    except Exception as e:
        print("STRIPE ERROR:", str(e))
        return jsonify({"error": str(e)}), 500

@app.route("/webhook", methods=["POST"])
def webhook():
    payload = request.data
    sig_header = request.headers.get("Stripe-Signature")
    webhook_secret = os.getenv("STRIPE_WEBHOOK_SECRET", "")

    try:
        if webhook_secret:
            event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
        else:
            event = stripe.Event.construct_from(request.json, stripe.api_key)
    except Exception as e:
        return jsonify({"error": str(e)}), 400

    if event["type"] == "checkout.session.completed":
        session = event["data"]["object"]
        user_id = session["metadata"]["user_id"]
        plan = session["metadata"]["plan"]
        db_patch("users", f"id=eq.{user_id}", {"plan": plan})

    return jsonify({"status": "ok"})

@app.route("/chat", methods=["POST"])
def chat_with_resume():
    data = request.json
    message = data.get("message", "")
    resume = data.get("resume", "")
    
    prompt = f"""
    You are a professional resume coach. The user has just received this resume and cover letter:
    
    ---
    {resume}
    ---
    
    The user is now asking you to help them refine it. Respond helpfully to their request.
    If they ask you to change something, output the FULL updated resume and cover letter.
    If they ask a question, just answer it conversationally.
    
    User message: {message}
    """
    
    chat = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )
    
    return jsonify({"result": chat.choices[0].message.content})

@app.route("/usage", methods=["POST"])
def get_usage():
    data = request.json
    user_id = data.get("user_id")
    today = str(date.today())
    usage = db_get("usage", f"user_id=eq.{user_id}&date=eq.{today}&select=count")
    count = usage[0]["count"] if usage else 0
    return jsonify({"count": count})

@app.route("/score", methods=["POST"])
def score_resume():
    data = request.json
    job_description = data.get("job_description", "")
    resume = data.get("resume", "")

    prompt = f"""
    You are an ATS (Applicant Tracking System) expert.
    
    Analyze this resume against the job description and respond ONLY with a JSON object like this:
    {{
      "score": 82,
      "matched_keywords": ["Python", "React", "Agile"],
      "missing_keywords": ["Docker", "Kubernetes"],
      "tip": "Add Docker experience to significantly boost your match score."
    }}

    Job Description:
    {job_description}

    Resume:
    {resume}

    Rules:
    - Score between 0-100
    - matched_keywords: important keywords found in both
    - missing_keywords: important keywords in job but missing from resume
    - tip: one specific actionable tip
    - Respond with ONLY the JSON, no extra text
    """

    chat = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )

    try:
        import json
        result = json.loads(chat.choices[0].message.content)
    except:
        result = {"score": 75, "matched_keywords": [], "missing_keywords": [], "tip": "Review your resume against the job description."}

    return jsonify(result)


@app.route("/linkedin", methods=["POST"])
def linkedin_bio():
    data = request.json
    plan = data.get("plan", "free")

    if plan != "proplus":
        return jsonify({"error": "Pro+ required"}), 403

    experience = data.get("experience", "")
    skills = data.get("skills", "")
    tone = data.get("tone", "professional")
    voice_sample = data.get("voice_sample", "")

    voice_instruction = f"""
    Match the writing style of this sample:
    ---
    {voice_sample}
    ---
    """ if voice_sample else ""

    prompt = f"""
    You are a LinkedIn profile expert.

    Write a compelling LinkedIn About section for this person:
    - Experience: {experience}
    - Skills: {skills}
    - Tone: {tone}

    {voice_instruction}

    Rules:
    - Maximum 300 words
    - First line must be a hook that grabs attention
    - Write in first person
    - Sound human, not like an AI
    - End with what they are looking for or open to
    - Do NOT use hashtags
    - Output ONLY the bio text, nothing else
    """

    chat = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )

    return jsonify({"result": chat.choices[0].message.content})


@app.route("/improve", methods=["POST"])
def improve_resume():
    data = request.json
    resume = data.get("resume", "")
    job_description = data.get("job_description", "")
    instruction = data.get("instruction", "Improve this resume")

    prompt = f"""
    You are an expert resume coach.

    {"Here is the target job description:" + job_description if job_description else ""}

    Here is the resume to improve:
    {resume}

    User instruction: {instruction}

    Rules:
    - Output the FULL improved resume
    - Keep the same structure but improve the content
    - If a job description is provided, tailor keywords to match it
    - Sound human and natural
    - Do NOT explain what you changed, just output the resume
    """

    chat = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )

    return jsonify({"result": chat.choices[0].message.content})

@app.route("/forgot-password", methods=["POST"])
def forgot_password():
    data = request.json
    email = data.get("email", "").lower().strip()

    users = db_get("users", f"email=eq.{email}&select=id")
    if not users:
        return jsonify({"message": "If this email exists you'll receive a reset link shortly!"}), 200

    token = secrets.token_urlsafe(32)
    db_post("password_resets", {"email": email, "token": token})

    reset_url = f"https://careercraft-rouge.vercel.app/index.html?reset={token}&email={email}"

    try:
        resend.Emails.send({
            "from": "CareerCraft <onboarding@resend.dev>",
            "to": email,
            "subject": "Reset your CareerCraft password",
            "html": f"""
            <div style="font-family:sans-serif;max-width:500px;margin:0 auto;padding:2rem">
              <h1 style="font-size:1.5rem;color:#1a1a2e">Reset your password</h1>
              <p style="color:#666;line-height:1.6">Click the button below to reset your CareerCraft password. This link expires in 1 hour.</p>
              <a href="{reset_url}" style="display:inline-block;margin:1.5rem 0;padding:12px 28px;background:#7c6aff;color:white;border-radius:8px;text-decoration:none;font-weight:600">Reset Password</a>
              <p style="color:#999;font-size:0.8rem">If you didn't request this, ignore this email.</p>
            </div>
            """
        })
    except Exception as e:
        print("EMAIL ERROR:", str(e))

    return jsonify({"message": "If this email exists you'll receive a reset link shortly!"})


@app.route("/reset-password", methods=["POST"])
def reset_password():
    data = request.json
    email = data.get("email", "").lower().strip()
    token = data.get("token", "")
    new_password = data.get("password", "")

    if len(new_password) < 6:
        return jsonify({"error": "Password must be at least 6 characters"}), 400

    resets = db_get("password_resets", f"email=eq.{email}&token=eq.{token}&select=id,created_at")
    if not resets:
        return jsonify({"error": "Invalid or expired reset link"}), 400

    hashed = bcrypt.generate_password_hash(new_password).decode("utf-8")
    db_patch("users", f"email=eq.{email}", {"password": hashed})

    requests.delete(
        f"{SUPABASE_URL}/rest/v1/password_resets?email=eq.{email}",
        headers=HEADERS
    )

    return jsonify({"message": "Password reset successfully!"})

@app.route("/resumes", methods=["POST"])
def get_resumes():
    data = request.json
    user_id = data.get("user_id")
    resumes = db_get("resumes", f"user_id=eq.{user_id}&select=id,job_title,created_at&order=created_at.desc&limit=20")
    return jsonify({"resumes": resumes})

@app.route("/resume", methods=["POST"])
def get_resume():
    data = request.json
    resume_id = data.get("resume_id")
    user_id = data.get("user_id")
    resume = db_get("resumes", f"id=eq.{resume_id}&user_id=eq.{user_id}&select=id,job_title,content,created_at")
    if not resume:
        return jsonify({"error": "Resume not found"}), 404
    return jsonify({"resume": resume[0]})


# ─── College Application Routes ────────────────────────────────────────────────

@app.route("/college/personal-statement", methods=["POST"])
def college_personal_statement():
    data = request.json
    user_id = data.get("user_id")
    plan = data.get("plan", "free")

    # Rate limit free users using the same usage table
    if plan == "free" and user_id:
        today = str(date.today())
        usage = db_get("usage", f"user_id=eq.{user_id}&date=eq.{today}&select=count,id")
        if usage:
            if usage[0]["count"] >= FREE_DAILY_LIMIT:
                return jsonify({"error": "daily_limit"}), 429
            db_patch("usage", f"user_id=eq.{user_id}&date=eq.{today}", {"count": usage[0]["count"] + 1})
        else:
            db_post("usage", {"user_id": user_id, "date": today, "count": 1})

    university = data.get("university", "")
    story = data.get("story", "")
    academics = data.get("academics", "")
    tone = data.get("tone", "Authentic & personal")
    prompt_q = data.get("prompt", "")
    word_limit = data.get("word_limit", 650)
    voice_sample = data.get("voice_sample", "")

    voice_instruction = ""
    if voice_sample and plan in ["pro", "proplus"]:
        voice_instruction = f"""
        The applicant has provided a writing sample — match their voice exactly:
        ---
        {voice_sample}
        ---
        Analyze their vocabulary, sentence rhythm, and personality. Write as them, not as a generic AI.
        """

    prompt = f"""
    You are an expert college admissions counselor with 20 years of experience helping students get into top universities.

    Write a compelling personal statement for a college application.

    Target university/program: {university or "a competitive university"}
    {"Essay prompt: " + prompt_q if prompt_q else "Open-ended personal statement"}
    Requested tone: {tone}
    Word limit: approximately {word_limit} words

    About the applicant:
    - Their story: {story}
    - Academic achievements & interests: {academics or "Not specified"}

    {voice_instruction}

    Rules:
    1. NEVER invent experiences, awards, or facts not mentioned by the applicant
    2. Stay within the word limit (±10%)
    3. Open with a compelling hook — NOT "I have always been passionate about..."
    4. Show, don't tell — use specific moments and details
    5. Sound like a real {tone.lower()} student, not a robot
    6. End with a strong forward-looking statement connecting past to future
    7. Output ONLY the personal statement — no intro, no explanation, no notes
    """

    chat = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )
    return jsonify({"result": chat.choices[0].message.content})


@app.route("/college/activities", methods=["POST"])
def college_activities():
    data = request.json
    activities = data.get("activities", [])
    university = data.get("university", "")

    if not activities:
        return jsonify({"error": "No activities provided"}), 400

    activities_text = "\n".join([
        f"{i+1}. {a.get('name','')} — {a.get('role','')} — {a.get('desc','')}"
        for i, a in enumerate(activities)
    ])

    prompt = f"""
    You are a college admissions expert helping a student write their extracurricular activity descriptions.

    {f"Target school: {university}" if university else ""}

    Here are the student's activities (raw notes):
    {activities_text}

    For each activity, write a polished 1-2 sentence description (maximum 150 characters per activity, like Common App format).
    Focus on:
    - Specific role and impact
    - Leadership, initiative, or growth shown
    - Quantify where possible (team size, hours/week, achievements)
    - Strong action verbs

    Format output as:
    [Activity Name] | [Role]
    [Polished description]

    Output ONLY the formatted activity list. No intro, no notes.
    """

    chat = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )
    return jsonify({"result": chat.choices[0].message.content})


@app.route("/college/short-answer", methods=["POST"])
def college_short_answer():
    data = request.json
    question = data.get("question", "")
    background = data.get("background", "")
    university = data.get("university", "")
    word_limit = data.get("word_limit", 150)
    tone = data.get("tone", "Genuine & direct")
    voice_sample = data.get("voice_sample", "")
    plan = data.get("plan", "free")

    voice_instruction = ""
    if voice_sample and plan in ["pro", "proplus"]:
        voice_instruction = f"""
        Match the writing voice of this sample exactly:
        ---
        {voice_sample}
        ---
        Analyze their vocabulary, rhythm, and personality. Write as them.
        """

    if not question or not background:
        return jsonify({"error": "Question and background are required"}), 400

    prompt = f"""
    You are a college admissions expert helping a student answer a short application question.

    Question: {question}
    {f"University/Program: {university}" if university else ""}
    Word limit: approximately {word_limit} words
    Tone: {tone}

    About the student: {background}

    {voice_instruction}

    Rules:
    1. Answer the question directly and specifically
    2. Stay within the word limit (±10%)
    3. NEVER invent facts not mentioned by the student
    4. Sound like a real human student — {tone.lower()}
    5. Be specific, not generic — avoid clichés
    6. Output ONLY the answer — no intro, no explanation
    """

    chat = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )
    return jsonify({"result": chat.choices[0].message.content})


@app.route("/college/summary", methods=["POST"])
def college_summary():
    data = request.json
    user_id = data.get("user_id")
    plan = data.get("plan", "free")

    if plan == "free" and user_id:
        today = str(date.today())
        usage = db_get("usage", f"user_id=eq.{user_id}&date=eq.{today}&select=count,id")
        if usage:
            if usage[0]["count"] >= FREE_DAILY_LIMIT:
                return jsonify({"error": "daily_limit"}), 429
            db_patch("usage", f"user_id=eq.{user_id}&date=eq.{today}", {"count": usage[0]["count"] + 1})
        else:
            db_post("usage", {"user_id": user_id, "date": today, "count": 1})

    university = data.get("university", "the university")
    program = data.get("program", "")
    story = data.get("story", "")
    academics = data.get("academics", "")
    activities = data.get("activities", "")
    why = data.get("why", "")
    goals = data.get("goals", "")
    voice_sample = data.get("voice_sample", "")

    voice_instruction = ""
    if voice_sample and plan in ["pro", "proplus"]:
        voice_instruction = f"""
        Match the writing voice of this sample exactly across all sections:
        ---
        {voice_sample}
        ---
        Analyze their vocabulary, sentence rhythm, and personality. Write as them, not as a generic AI.
        """
    voice_sample = data.get("voice_sample", "")

    voice_instruction = ""
    if voice_sample and plan in ["pro", "proplus"]:
        voice_instruction = f"""
        Match the writing voice of this sample exactly across all sections:
        ---
        {voice_sample}
        ---
        Analyze their vocabulary, sentence rhythm, and personality. Write everything as them.
        """

    import json as json_lib

    prompt = f"""
    You are an expert college admissions counselor. Generate a full application package for this student.

    Target: {university} — {program}

    Student background:
    - Story: {story}
    - Academics: {academics}
    - Activities: {activities}
    - Why this school: {why}
    - Goals: {goals}

    {voice_instruction}

    Generate three sections and respond ONLY with a JSON object:
    {{
      "personal_statement": "~500 word personal statement...",
      "activities_summary": "Polished activity overview paragraph (150-200 words)...",
      "why_us": "Compelling why this school response (150-200 words)..."
    }}

    Rules for all sections:
    - Never invent facts not provided
    - Sound human and authentic
    - Be specific, not generic
    - Respond with ONLY the JSON, no extra text or markdown
    """

    chat = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )

    try:
        content = chat.choices[0].message.content
        # Strip potential markdown code fences
        content = content.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
        result = json_lib.loads(content)
    except Exception:
        result = {
            "personal_statement": chat.choices[0].message.content,
            "activities_summary": "Could not parse activities section. Please regenerate.",
            "why_us": "Could not parse why-us section. Please regenerate."
        }

    return jsonify(result)
@app.route("/scrape-job", methods=["POST"])
def scrape_job():
    import re
    from html.parser import HTMLParser

    data = request.json
    url = data.get("url", "").strip()

    if not url or not url.startswith("http"):
        return jsonify({"error": "Invalid URL"}), 400

    HEADERS_SCRAPE = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xhtml+xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
    }

    try:
        resp = requests.get(url, headers=HEADERS_SCRAPE, timeout=12)
        resp.raise_for_status()
        html = resp.text
    except Exception as e:
        return jsonify({"error": f"Could not fetch URL: {str(e)}"}), 400

    # Strip scripts, styles, nav, footer
    html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r'<nav[^>]*>.*?</nav>', '', html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r'<footer[^>]*>.*?</footer>', '', html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r'<header[^>]*>.*?</header>', '', html, flags=re.DOTALL | re.IGNORECASE)

    # Convert <br>, <p>, <li> to newlines for readability
    html = re.sub(r'<br\s*/?>', '\n', html, flags=re.IGNORECASE)
    html = re.sub(r'</?p[^>]*>', '\n', html, flags=re.IGNORECASE)
    html = re.sub(r'<li[^>]*>', '\n• ', html, flags=re.IGNORECASE)

    # Strip remaining tags
    clean = re.sub(r'<[^>]+>', ' ', html)
    # Collapse whitespace
    clean = re.sub(r'[ \t]+', ' ', clean)
    clean = re.sub(r'\n{3,}', '\n\n', clean)
    clean = clean.strip()

    if len(clean) < 200:
        return jsonify({"error": "Could not extract enough content from this page"}), 400

    # Detect hostname for source label
    try:
        from urllib.parse import urlparse
        hostname = urlparse(url).hostname or url
        if 'linkedin' in hostname: source = 'LinkedIn'
        elif 'indeed' in hostname: source = 'Indeed'
        elif 'glassdoor' in hostname: source = 'Glassdoor'
        elif 'greenhouse' in hostname: source = 'Greenhouse'
        elif 'lever' in hostname: source = 'Lever'
        elif 'workday' in hostname: source = 'Workday'
        else: source = hostname
    except:
        source = 'job posting'

    # Use Groq to extract just the job description cleanly
    prompt = f"""
    The following is raw scraped text from a job posting page. Extract ONLY the job description content — 
    job title, company, responsibilities, requirements, and qualifications.
    Remove any navigation, cookie notices, ads, or unrelated site content.
    Format it cleanly with clear sections. Output ONLY the cleaned job description, nothing else.

    Raw text:
    {clean[:6000]}
    """

    try:
        chat = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
        )
        job_description = chat.choices[0].message.content.strip()
    except Exception:
        # Fallback: just return first 3000 chars of cleaned text
        job_description = clean[:3000]

    return jsonify({"job_description": job_description, "source": source})


@app.route("/college/research-uni", methods=["POST"])
def college_research_uni():
    data = request.json
    university = data.get("university", "").strip()
    if not university:
        return jsonify({"research": "", "summary": "No university provided."}), 400

    research_prompt = f"""What are the application requirements for {university}?

Include ALL of the following you know about:
1. Application system (Common App, UCAS, Coalition, own portal, etc.)
2. Personal statement / main essay — prompt(s) and word limit
3. Supplemental essays — exact prompts and word limits
4. Activities/extracurriculars section format
5. Any "Why us?" or program-specific questions
6. Any unique requirements for this school

Be as specific as possible with actual prompts and word counts. Format as clear bullet points."""

    chat = client.chat.completions.create(
        messages=[{"role": "user", "content": research_prompt}],
        model="llama-3.3-70b-versatile",
        max_tokens=700,
    )
    research = chat.choices[0].message.content.strip()

    summary_chat = client.chat.completions.create(
        messages=[{"role": "user", "content": f"Summarize in 2-3 short sentences for a student — the most important things to know about applying to {university}:\n\n{research}"}],
        model="llama-3.3-70b-versatile",
        max_tokens=120,
    )
    summary = summary_chat.choices[0].message.content.strip()

    return jsonify({"research": research, "summary": summary})

@app.route("/college/export-docx", methods=["POST"])
def college_export_docx():
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    data = request.json
    title = data.get("title", "Application")
    content = data.get("content", "")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.runs[0].font.size = Pt(16)
    doc.add_paragraph()

    for block in content.split('\n\n'):
        block = block.strip()
        if not block:
            continue
        if block.startswith('---') or (block.isupper() and len(block) < 60):
            doc.add_heading(block.strip('-').strip(), level=2)
        else:
            p = doc.add_paragraph(block)
            p.paragraph_format.space_after = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name=f"{title}.docx")


@app.route("/college/chat", methods=["POST"])
def college_chat():
    data = request.json
    mode = data.get("mode", "summary")
    history = data.get("history", [])
    plan = data.get("plan", "free")
    uni_research = data.get("uni_research", "")
    is_pro = plan in ["pro", "proplus"]

    research_context = f"""
UNIVERSITY REQUIREMENTS (researched):
{uni_research}

Use this to ask questions specific to what this university actually needs — 
e.g. if they require a "Why Northwestern?" supplement, ask about that specifically.
If they use UCAS, follow UCAS rules. If Common App prompt 2, ask about that prompt.
""" if uni_research else ""

    system = f"""You are an expert college admissions counselor helping a student write a full application package (personal statement + activities + why us + any supplementals).

Have a warm natural conversation to gather everything needed. Rules:
1. Ask ONE question at a time
2. Use the university requirements below to ask the RIGHT questions for THIS specific school
3. Adapt based on actual answers — dig into specifics they mention
4. After 6-8 exchanges you likely have enough
5. Near the end, if is_pro={is_pro}, ask for voice sample: start message with "VOICE_STEP: "
6. When ready: respond with exactly "READY_TO_GENERATE"
7. Output ONLY the question, "VOICE_STEP: [question]", or "READY_TO_GENERATE"

{research_context}"""

    messages = []
    for h in history:
        messages.append({{"role": h["role"], "content": h["content"]}})

    if not messages:
        messages.append({{"role": "user", "content": "I want to write my full college application package. Start by asking me what university I'm applying to."}})

    chat = client.chat.completions.create(
        messages=[{{"role": "system", "content": system}}] + messages,
        model="llama-3.3-70b-versatile",
        max_tokens=300,
    )

    response = chat.choices[0].message.content.strip()

    if response.startswith("VOICE_STEP:"):
        return jsonify({{"message": response.replace("VOICE_STEP:", "").strip(), "voice_step": True, "done": False}})
    elif "READY_TO_GENERATE" in response:
        closing = response.replace("READY_TO_GENERATE", "").strip() or "Perfect — writing your full application package now ✦"
        return jsonify({{"message": closing, "done": True}})
    else:
        return jsonify({{"message": response, "done": False, "placeholder": ""}})

    mode_label = {
        "personal": "personal statement / college essay",
        "activities": "extracurricular activity list",
        "shortanswer": "short answer supplemental essay",
        "summary": "full application package"
    }.get(mode, "college application")

    is_pro = plan in ["pro", "proplus"]

    system = f"""You are an expert college admissions counselor helping a student write their {mode_label}.

Have a warm, natural conversation to gather everything needed. Rules:
1. Ask ONE question at a time — never multiple questions in one message
2. ADAPT based on answers — if they say MIT ask about research; UCAS means 4000 chars; Common App ask which of 5 prompts; Oxford/Cambridge means deep intellectual curiosity focus; UC schools means PIQs of 350 words; Canadian unis often have shorter supplementals
3. Ask specific follow-up questions based on what they actually share
4. After 6-8 good exchanges you likely have enough info
5. Near the end (but not first), if is_pro={is_pro}, ask for a voice sample by starting your message with exactly "VOICE_STEP: " followed by the question
6. When you have enough, respond with exactly: READY_TO_GENERATE
7. Output ONLY the question text, or "VOICE_STEP: [question]", or "READY_TO_GENERATE"

University knowledge:
- Common App: 650 words, 5 prompts
- UCAS: 4000 chars, one statement, academic focus
- Oxford/Cambridge: UCAS but needs deep intellectual curiosity
- UC schools: 4 PIQs x 350 words
- Canadian unis: shorter supplementals, program-specific questions
- MIT/Ivy League: research, intellectual curiosity, unique angle
- Australian unis: often grades-based, sometimes a personal statement"""

    messages = []
    for h in history:
        messages.append({"role": h["role"], "content": h["content"]})

    if not messages:
        messages.append({"role": "user", "content": f"I want to write my {mode_label}. Start the conversation."})

    chat = client.chat.completions.create(
        messages=[{"role": "system", "content": system}] + messages,
        model="llama-3.3-70b-versatile",
        max_tokens=300,
    )

    response = chat.choices[0].message.content.strip()

    if response.startswith("VOICE_STEP:"):
        msg = response.replace("VOICE_STEP:", "").strip()
        return jsonify({"message": msg, "voice_step": True, "done": False})
    elif "READY_TO_GENERATE" in response:
        closing = response.replace("READY_TO_GENERATE", "").strip() or "Perfect — I have everything I need. Writing your application now ✦"
        return jsonify({"message": closing, "done": True})
    else:
        return jsonify({"message": response, "done": False, "placeholder": ""})


@app.route("/college/generate", methods=["POST"])
def college_generate():
    import json as json_lib
    data = request.json
    history = data.get("history", [])
    voice_sample = data.get("voice_sample")
    plan = data.get("plan", "free")
    user_id = data.get("user_id")
    uni_research = data.get("uni_research", "")

    if plan == "free" and user_id:
        today = str(date.today())
        usage = db_get("usage", f"user_id=eq.{user_id}&date=eq.{today}&select=count,id")
        if usage:
            if usage[0]["count"] >= FREE_DAILY_LIMIT:
                return jsonify({"error": "daily_limit"}), 429
            db_patch("usage", f"user_id=eq.{user_id}&date=eq.{today}", {"count": usage[0]["count"] + 1})
        else:
            db_post("usage", {"user_id": user_id, "date": today, "count": 1})

    conversation = "\n".join([f"{h['role'].upper()}: {h['content']}" for h in history])

    voice_instruction = ""
    if voice_sample and plan in ["pro", "proplus"]:
        voice_instruction = f"""VOICE MATCHING — CRITICAL: Write everything in the student's exact voice:
---
{voice_sample}
---"""

    research_context = f"""
UNIVERSITY REQUIREMENTS (researched):
{uni_research}

Apply these EXACTLY — use the correct word limits, answer the actual prompts, 
follow the right application format (UCAS/Common App/etc.), and include any required supplementals.
""" if uni_research else ""

    prompt = f"""You are an expert college admissions counselor. Write a complete application package based on this conversation.

CONVERSATION:
{conversation}

{research_context}
{voice_instruction}

Respond ONLY with this JSON (no markdown, no explanation):
{{
  "personal_statement": "Full personal statement here — use the word limit and prompt from the conversation/research...",
  "activities_summary": "Polished activities overview 150-200 words...",
  "why_us": "Compelling why this school 150-200 words...",
  "extras": "Any required supplemental essays or additional sections specific to this university (leave empty string if none needed)"
}}

Rules:
- NEVER invent facts not mentioned in the conversation
- Sound like a real human — open with a specific moment or image, not "I have always been..."
- Apply the EXACT prompts, word limits, and format from the university requirements
- If UCAS: personal statement is 4000 chars max, no "why us" section needed
- If UC PIQs: write as 4 separate responses of 350 words each in the personal_statement field
- If Common App: follow the specific prompt chosen
- If extras are required by this uni (e.g. Northwestern supplement, MIT essays): write them in the extras field"""

    chat = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="llama-3.3-70b-versatile",
    )
    raw = chat.choices[0].message.content.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
    try:
        result = json_lib.loads(raw)
    except:
        result = {"personal_statement": raw, "activities_summary": "", "why_us": "", "extras": ""}
    return jsonify(result)

    if plan == "free" and user_id:
        today = str(date.today())
        usage = db_get("usage", f"user_id=eq.{user_id}&date=eq.{today}&select=count,id")
        if usage:
            if usage[0]["count"] >= FREE_DAILY_LIMIT:
                return jsonify({"error": "daily_limit"}), 429
            db_patch("usage", f"user_id=eq.{user_id}&date=eq.{today}", {"count": usage[0]["count"] + 1})
        else:
            db_post("usage", {"user_id": user_id, "date": today, "count": 1})

    conversation = "\n".join([f"{h['role'].upper()}: {h['content']}" for h in history])

    voice_instruction = ""
    if voice_sample and plan in ["pro", "proplus"]:
        voice_instruction = f"""VOICE MATCHING: Write in the student's exact voice based on this sample:
---
{voice_sample}
---"""

    if mode == "summary":
        prompt = f"""Expert college admissions counselor. Write a full application package based on this conversation.

CONVERSATION:
{conversation}

{voice_instruction}

Respond ONLY with this JSON (no markdown):
{{"personal_statement": "500-650 words...", "activities_summary": "150-200 words...", "why_us": "150-200 words..."}}

Rules: Never invent facts. Sound human. Apply university-specific requirements from conversation."""

        chat = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
        )
        raw = chat.choices[0].message.content.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
        try:
            result = json_lib.loads(raw)
        except:
            result = {"personal_statement": raw, "activities_summary": "", "why_us": ""}
        return jsonify(result)

    else:
        task = {
            "personal": "Write a compelling personal statement/college essay",
            "activities": "Write polished activity descriptions (Common App format: ~150 chars each, action verbs)",
            "shortanswer": "Write a focused short answer response"
        }.get(mode, "Write the college application content")

        prompt = f"""Expert college admissions counselor. Based on this conversation, {task}.

CONVERSATION:
{conversation}

{voice_instruction}

Rules:
- Never invent facts not in conversation
- Apply exact word/char limit discussed
- Open with a hook, not "I have always been..."
- Be specific to this student's actual experiences and university
- Apply uni-specific format (UCAS 4000 chars, UC PIQs, Common App prompts, etc.)
- Output ONLY the final content — no intro or explanation"""

        chat = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama-3.3-70b-versatile",
        )
        return jsonify({"result": chat.choices[0].message.content.strip()})


if __name__ == "__main__":
    app.run(debug=True)
